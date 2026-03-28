"""
File-based adapters: PDF, Markdown/TXT, DOCX, XLSX.
"""
import os
import logging
from typing import List

from .base import SourceAdapter, IngestionResult, SourceType

logger = logging.getLogger(__name__)


class PdfAdapter(SourceAdapter):
    """Extract text from PDF files using PyMuPDF (fitz)."""

    def can_handle(self, source_ref: str) -> bool:
        return source_ref.lower().endswith('.pdf')

    def ingest(self, source_ref: str, **kwargs) -> IngestionResult:
        import fitz  # PyMuPDF
        doc = fitz.open(source_ref)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()

        full_text = "\n\n".join(pages)
        return IngestionResult(
            text=full_text,
            metadata={
                "source": source_ref,
                "format": "pdf",
                "page_count": len(pages),
            },
            source_type=SourceType.FILE,
        )


class TextAdapter(SourceAdapter):
    """Extract text from plain text / Markdown files."""

    EXTENSIONS = ('.txt', '.md', '.markdown', '.rst', '.log')

    def can_handle(self, source_ref: str) -> bool:
        return source_ref.lower().endswith(self.EXTENSIONS)

    def ingest(self, source_ref: str, **kwargs) -> IngestionResult:
        encoding = kwargs.get('encoding', 'utf-8')
        with open(source_ref, 'r', encoding=encoding) as f:
            text = f.read()
        return IngestionResult(
            text=text,
            metadata={"source": source_ref, "format": os.path.splitext(source_ref)[1].lstrip('.')},
            source_type=SourceType.FILE,
        )


class DocxAdapter(SourceAdapter):
    """Extract text from DOCX files using python-docx."""

    def can_handle(self, source_ref: str) -> bool:
        return source_ref.lower().endswith(('.docx', '.doc'))

    def ingest(self, source_ref: str, **kwargs) -> IngestionResult:
        from docx import Document
        doc = Document(source_ref)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables_text: List[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables_text.append(" | ".join(cells))

        full_text = "\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n=== Tables ===\n" + "\n".join(tables_text)

        return IngestionResult(
            text=full_text,
            metadata={"source": source_ref, "format": "docx"},
            source_type=SourceType.FILE,
        )


class ExcelAdapter(SourceAdapter):
    """Extract data from Excel (XLSX/XLS) and convert to natural-language text."""

    def can_handle(self, source_ref: str) -> bool:
        return source_ref.lower().endswith(('.xlsx', '.xls'))

    def ingest(self, source_ref: str, **kwargs) -> IngestionResult:
        import pandas as pd

        sheet_name = kwargs.get('sheet_name', None)  # None = all sheets
        xls = pd.ExcelFile(source_ref)
        sheets = [sheet_name] if sheet_name else xls.sheet_names

        all_text_parts: List[str] = []
        total_rows = 0

        for sheet in sheets:
            df = pd.read_excel(source_ref, sheet_name=sheet)
            total_rows += len(df)
            schema = f"Sheet '{sheet}': {len(df)} rows, columns: {', '.join(df.columns)}"
            all_text_parts.append(schema)

            row_limit = kwargs.get('row_limit', 500)
            for _, row in df.head(row_limit).iterrows():
                parts = [f"{col} is {val}" for col, val in row.items() if pd.notna(val)]
                all_text_parts.append(". ".join(parts) + ".")

        return IngestionResult(
            text="\n".join(all_text_parts),
            metadata={
                "source": source_ref,
                "format": "excel",
                "sheets": sheets,
                "total_rows": total_rows,
            },
            source_type=SourceType.FILE,
        )


class HtmlAdapter(SourceAdapter):
    """Extract readable text from HTML files using BeautifulSoup."""

    def can_handle(self, source_ref: str) -> bool:
        return source_ref.lower().endswith(('.html', '.htm'))

    def ingest(self, source_ref: str, **kwargs) -> IngestionResult:
        from bs4 import BeautifulSoup

        with open(source_ref, 'r', encoding=kwargs.get('encoding', 'utf-8')) as f:
            soup = BeautifulSoup(f, 'html.parser')

        # Remove script/style elements
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()

        text = soup.get_text(separator='\n', strip=True)
        title = soup.title.string if soup.title else os.path.basename(source_ref)

        return IngestionResult(
            text=text,
            metadata={"source": source_ref, "format": "html", "title": title},
            source_type=SourceType.FILE,
        )
